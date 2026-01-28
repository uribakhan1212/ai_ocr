import streamlit as st
import google.generativeai as genai
from docx import Document
from PIL import Image
import io
import os
from datetime import datetime

# Configure page
st.set_page_config(
    page_title="Image Text Extractor",
    page_icon="📄",
    layout="wide"
)

def configure_gemini():
    """Configure Gemini API with user's API key"""
    api_key = st.session_state.get('api_key', '')
    if api_key:
        genai.configure(api_key=api_key)
        return True
    return False

def extract_text_from_image(image, prompt_type="comprehensive"):
    """Extract text from image using Gemini Vision API"""
    try:
        # Use Gemini 1.5 Flash model
        model = genai.GenerativeModel('models/gemini-2.5-flash')
        
        prompts = {
            "comprehensive": """
            Please extract ALL text content from this image. Include:
            1. All visible text (headings, paragraphs, captions, labels)
            2. Any structured data (tables, lists, forms)
            3. Text in different fonts, sizes, or styles
            4. Preserve the original formatting and structure as much as possible
            5. Include any numbers, dates, or special characters
            
            Format the output in a clear, readable way that maintains the document's structure.
            """,
            "ocr_only": "Extract only the text content from this image, preserving line breaks and structure.",
            "structured": "Extract text from this image and organize it with clear headings and sections."
        }
        
        response = model.generate_content([prompts[prompt_type], image])
        return response.text
    except Exception as e:
        st.error(f"Error extracting text: {str(e)}")
        return None

def create_word_document(extracted_text, filename="extracted_text.docx"):
    """Create a Word document with the extracted text, parsing markdown formatting"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Extracted Text Document', 0)
    
    # Add timestamp
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")  # Empty line
    
    # Add extracted content header
    doc.add_heading('Extracted Content', level=1)
    
    # Parse and format the extracted text
    lines = extracted_text.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        # Check if this is the start of a table
        if line.startswith('|') and '|' in line[1:]:
            # Process table
            table_lines = []
            j = i
            
            # Collect all table lines
            while j < len(lines) and lines[j].strip().startswith('|'):
                table_line = lines[j].strip()
                # Skip separator lines (like | :--- | :--- |)
                if not all(c in '|:-= ' for c in table_line):
                    table_lines.append(table_line)
                j += 1
            
            if table_lines:
                create_word_table(doc, table_lines)
            
            i = j
            continue
            
        # Handle different markdown elements
        if line.startswith('###'):
            # Level 3 heading
            heading_text = line.replace('###', '').strip()
            doc.add_heading(heading_text, level=3)
            
        elif line.startswith('##'):
            # Level 2 heading
            heading_text = line.replace('##', '').strip()
            doc.add_heading(heading_text, level=2)
            
        elif line.startswith('#'):
            # Level 1 heading
            heading_text = line.replace('#', '').strip()
            doc.add_heading(heading_text, level=1)
            
        elif line.startswith('*') or line.startswith('-'):
            # Bullet point
            bullet_text = line[1:].strip()
            # Parse bold text in bullet points
            bullet_text = parse_inline_formatting(bullet_text)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, bullet_text)
            
        elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            # Numbered list
            num_text = line.split('.', 1)[1].strip() if '.' in line else line
            num_text = parse_inline_formatting(num_text)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, num_text)
            
        else:
            # Regular paragraph
            if line:
                formatted_text = parse_inline_formatting(line)
                p = doc.add_paragraph()
                add_formatted_text(p, formatted_text)
        
        i += 1
    
    # Save to bytes buffer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer

def create_word_table(doc, table_lines):
    """Create a Word table from markdown table lines"""
    if not table_lines:
        return
    
    # Parse table data
    table_data = []
    for line in table_lines:
        # Split by | and clean up
        cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove first and last empty elements
        if cells:  # Only add non-empty rows
            table_data.append(cells)
    
    if not table_data:
        return
    
    # Create Word table
    rows = len(table_data)
    cols = len(table_data[0]) if table_data else 0
    
    if rows == 0 or cols == 0:
        return
    
    # Create table
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    # Fill table data
    for row_idx, row_data in enumerate(table_data):
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(table.rows[row_idx].cells):
                cell = table.rows[row_idx].cells[col_idx]
                
                # Parse formatting for cell content
                formatted_text = parse_inline_formatting(cell_data)
                
                # Clear existing content
                cell.text = ''
                
                # Add formatted content
                paragraph = cell.paragraphs[0]
                add_formatted_text(paragraph, formatted_text)
                
                # Make header row bold
                if row_idx == 0:
                    for run in paragraph.runs:
                        run.bold = True
    
    # Add some spacing after table
    doc.add_paragraph("")

def parse_inline_formatting(text):
    """Parse inline markdown formatting like **bold** and return formatted segments"""
    segments = []
    current_pos = 0
    
    while current_pos < len(text):
        # Find next bold marker
        bold_start = text.find('**', current_pos)
        
        if bold_start == -1:
            # No more bold text, add remaining as normal
            if current_pos < len(text):
                segments.append(('normal', text[current_pos:]))
            break
            
        # Add text before bold as normal
        if bold_start > current_pos:
            segments.append(('normal', text[current_pos:bold_start]))
        
        # Find end of bold text
        bold_end = text.find('**', bold_start + 2)
        if bold_end == -1:
            # No closing **, treat as normal text
            segments.append(('normal', text[bold_start:]))
            break
            
        # Add bold text
        bold_text = text[bold_start + 2:bold_end]
        segments.append(('bold', bold_text))
        
        current_pos = bold_end + 2
    
    return segments

def add_formatted_text(paragraph, formatted_segments):
    """Add formatted text segments to a paragraph"""
    if isinstance(formatted_segments, str):
        # Simple string, add as normal text
        paragraph.add_run(formatted_segments)
        return
        
    for format_type, text in formatted_segments:
        run = paragraph.add_run(text)
        if format_type == 'bold':
            run.bold = True

def main():
    st.title("📄 Image Text Extractor")
    st.markdown("Upload an image and extract all text content into an editable Word document using Google Gemini AI.")
    
    # Sidebar for API configuration
    with st.sidebar:
        st.header("🔧 Configuration")
        
        # API Key input
        api_key = st.text_input(
            "Google Gemini API Key",
            type="password",
            value=st.session_state.get('api_key', ''),
            help="Get your API key from Google AI Studio"
        )
        
        if api_key:
            st.session_state.api_key = api_key
            if configure_gemini():
                st.success("✅ API Key configured")
                
                # Show available models for debugging
                if st.checkbox("Show available models"):
                    try:
                        models = genai.list_models()
                        vision_models = [m for m in models if 'generateContent' in m.supported_generation_methods]
                        st.write("Available vision models:")
                        for model in vision_models[:5]:  # Show first 5
                            st.write(f"- {model.name}")
                    except Exception as e:
                        st.write(f"Error listing models: {e}")
            else:
                st.error("❌ Invalid API Key")
        
        st.markdown("---")
        
        # Extraction options
        st.subheader("Extraction Options")
        prompt_type = st.selectbox(
            "Extraction Mode",
            ["comprehensive", "ocr_only", "structured"],
            help="Choose how detailed the text extraction should be"
        )
    
    # Main content area
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📤 Upload Image")
        
        uploaded_file = st.file_uploader(
            "Choose an image file",
            type=['png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp'],
            help="Upload an image containing text to extract"
        )
        
        if uploaded_file is not None:
            # Display uploaded image
            image = Image.open(uploaded_file)
            st.image(image, caption="Uploaded Image", use_column_width=True)
            
            # Image info
            st.info(f"📊 Image size: {image.size[0]} x {image.size[1]} pixels")
    
    with col2:
        st.header("📝 Extracted Text")
        
        if uploaded_file is not None and st.session_state.get('api_key'):
            if st.button("🚀 Extract Text", type="primary"):
                with st.spinner("Extracting text from image..."):
                    # Extract text
                    extracted_text = extract_text_from_image(image, prompt_type)
                    
                    if extracted_text:
                        st.session_state.extracted_text = extracted_text
                        st.success("✅ Text extracted successfully!")
                    else:
                        st.error("❌ Failed to extract text")
        
        # Display extracted text if available
        if 'extracted_text' in st.session_state:
            st.subheader("Extracted Content")
            
            # Editable text area
            edited_text = st.text_area(
                "Edit the extracted text:",
                value=st.session_state.extracted_text,
                height=400,
                help="You can edit the extracted text before downloading"
            )
            
            # Download options
            st.subheader("📥 Download Options")
            
            col_download1, col_download2 = st.columns(2)
            
            with col_download1:
                # Download as Word document
                if st.button("📄 Download as Word"):
                    doc_buffer = create_word_document(edited_text)
                    st.download_button(
                        label="💾 Download DOCX",
                        data=doc_buffer.getvalue(),
                        file_name=f"extracted_text_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            
            with col_download2:
                # Download as plain text
                st.download_button(
                    label="📝 Download as TXT",
                    data=edited_text,
                    file_name=f"extracted_text_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain"
                )
    
    # Instructions
    if not st.session_state.get('api_key'):
        st.warning("⚠️ Please enter your Google Gemini API key in the sidebar to get started.")
        
        with st.expander("🔑 How to get a Gemini API Key"):
            st.markdown("""
            1. Go to [Google AI Studio](https://makersuite.google.com/app/apikey)
            2. Sign in with your Google account
            3. Click "Create API Key"
            4. Copy the generated API key
            5. Paste it in the sidebar configuration
            """)

if __name__ == "__main__":
    main()